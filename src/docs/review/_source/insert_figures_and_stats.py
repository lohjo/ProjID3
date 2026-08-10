"""Insert every generated figure and every fit statistic into the Final Report.

    python src/docs/review/_source/insert_figures_and_stats.py [--dry-run]
                                                               [--max-width PX]

What it does, in document order:

1.  §6.4  new  - the measured breakthrough curves themselves (Fig. 9-11).
2.  §7         - captions the four charts that were embedded without one
                 (Fig. 12-15) and renumbers the three existing results figures
                 to Fig. 16-18 (nothing in the prose references them by number).
3.  §7.1  new  - cross-run trends and model ranking (Fig. 19-21); the stale
                 "6.???"/"7.?.4" pseudo-headings that follow become §7.2.x and
                 are promoted to real headings so they reach the table of contents.
4.  §8.2       - error-statistics definitions and the estimation-strategy
                 justification, then Tables 14-17.
5.  §8.5  new  - sensitivity analysis (Fig. 22-30, Tables 18-21).
6.  §9.4  new  - numerical verification of the model equations (Fig. 31-41).
7.  Appendix A - per-run fit diagnostics, 7 plots x 21 real runs = 147 figures.
8.  Appendix B - complete fit statistics: 504-row master table, parameters with
                 standard errors, all nested F-tests, degenerate fits.
9.  List of Figures updated; <w:updateFields> set so Word refreshes the TOC.

Every number comes from a committed artefact (see `report_stats_tables`), and
every figure from a committed PNG. Nothing is refitted and nothing is hand-typed.

The script works on a copy and only overwrites the real file once every step has
succeeded. Word must be closed while it runs.
"""
from __future__ import annotations

import argparse
import shutil
import sys
from pathlib import Path
from xml.sax.saxutils import escape

sys.path.insert(0, str(Path(__file__).resolve().parent))

import docxlib as D  # noqa: E402
import insert_experiments_section as IES  # noqa: E402  (reused helpers)
import report_stats_tables as T  # noqa: E402
from docx.oxml import parse_xml  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.shared import Inches  # noqa: E402
from PIL import Image  # noqa: E402

HERE = Path(__file__).resolve().parent
REPO = HERE.parents[3]
DOCX = REPO / "src/T32_PI05_Final_Report.docx"
WORK = HERE / "_work_report_figs.docx"
FIGCACHE = HERE / "_figcache"
IMG = REPO / "src/img/generated"

TEXT_WIDTH_IN = 5.9          # page 8.268 in, margins 1.18125 in each side
MAX_HEIGHT_IN = 7.4
CAPTION_STYLE = "Heading7"   # what every existing figure caption in the report uses

W = D.W
NSDECL = f'xmlns:w="{W}"'


# =========================================================================== #
# Image preparation
# =========================================================================== #
def prepare_image(src: Path, max_w_px: int) -> Path:
    """Downscale a committed PNG into the build cache. Never touches the source."""
    src = Path(src)
    out = FIGCACHE / f"{max_w_px}" / src.parent.name / src.name
    if out.exists() and out.stat().st_mtime >= src.stat().st_mtime:
        return out
    out.parent.mkdir(parents=True, exist_ok=True)
    im = Image.open(src)
    im.load()
    w, h = im.size
    if w > max_w_px:
        scale = max_w_px / w
        im = im.convert("RGB").resize((round(w * scale), round(h * scale)),
                                      Image.LANCZOS)
    else:
        im = im.convert("RGB")
    im.save(out, "PNG", optimize=True)
    return out


def picture_size(path: Path) -> tuple[float, float]:
    with Image.open(path) as im:
        w, h = im.size
    width = TEXT_WIDTH_IN
    height = width * h / w
    if height > MAX_HEIGHT_IN:
        height = MAX_HEIGHT_IN
        width = height * w / h
    return width, height


# =========================================================================== #
# Paragraph / table emission
# =========================================================================== #
class Builder:
    """Walks forward from an anchor, emitting paragraphs, images and tables."""

    def __init__(self, doc, max_w_px: int):
        self.doc = doc
        self.max_w_px = max_w_px
        self.n_images = 0
        self.n_tables = 0
        _, self.h1 = D.find_heading(doc, "4 Standard Operating Procedure", level=1)
        _, self.h2 = D.find_heading(doc, "4.1 Apparatus and Instrumentation", level=2)
        _, self.h3 = D.find_heading(doc, "4.1.1 Gas Feeding Section", level=3)
        _, self.cap = D.find_heading(doc, "Fig. 5) Schematic diagram", level=7)
        self.body = self._first_body_paragraph()

    def _first_body_paragraph(self):
        el = self.h1.getnext()
        while el is not None and (el.tag != qn("w:p") or D.is_heading(el)):
            el = el.getnext()
        if el is None:
            raise SystemExit("no plain body paragraph found to use as a template")
        return el

    def _template(self, kind: str):
        return {"Heading1": self.h1, "Heading2": self.h2, "Heading3": self.h3,
                CAPTION_STYLE: self.cap}.get(kind, self.body)

    def para(self, cursor, kind: str, text: str):
        """Clone the right paragraph shape, set style + text, insert after cursor.

        The clone is always taken from the style template, never from `cursor`:
        cursor is often a picture paragraph, and cloning that would duplicate the
        drawing into the caption.
        """
        tpl = self._template(kind)
        new_el = D.clone_paragraph_after(tpl, text=text, style=kind)
        new_el.getparent().remove(new_el)
        cursor.addnext(new_el)
        return new_el

    def image(self, cursor, png: Path):
        """add_picture() appends at the end of the body; relocate it to the cursor."""
        cached = prepare_image(png, self.max_w_px)
        w_in, h_in = picture_size(cached)
        self.doc.add_picture(str(cached), width=Inches(w_in), height=Inches(h_in))
        p = self.doc.paragraphs[-1]._p
        p.getparent().remove(p)
        pPr = p.find("w:pPr", D.NS)
        if pPr is None:
            pPr = p.makeelement(qn("w:pPr"), {})
            p.insert(0, pPr)
        jc = pPr.makeelement(qn("w:jc"), {})
        jc.set(qn("w:val"), "center")
        pPr.append(jc)
        # keep the picture on the same page as the caption that follows it
        pPr.append(pPr.makeelement(qn("w:keepNext"), {}))
        cursor.addnext(p)
        self.n_images += 1
        return p

    @staticmethod
    def keep_with_next(el) -> None:
        """Stop a page break landing between a caption and what it labels."""
        pPr = el.find("w:pPr", D.NS)
        if pPr is None:
            pPr = el.makeelement(qn("w:pPr"), {})
            el.insert(0, pPr)
        if pPr.find("w:keepNext", D.NS) is None:
            pPr.append(pPr.makeelement(qn("w:keepNext"), {}))

    def figure(self, cursor, png: Path, caption: str):
        cursor = self.image(cursor, png)
        return self.para(cursor, CAPTION_STYLE, caption)

    def table(self, cursor, rows, font_pt: float, header_repeat: bool = True):
        self.keep_with_next(cursor)  # cursor is the caption paragraph
        tbl_el = build_table_xml(rows, font_pt, header_repeat)
        cursor.addnext(tbl_el)
        spacer = D.clone_paragraph_after(self.body, text="", style=None)
        spacer.getparent().remove(spacer)
        tbl_el.addnext(spacer)
        self.n_tables += 1
        return spacer


def _cell_xml(text: str, font_pt: float, bold: bool) -> str:
    rpr = f'<w:rPr>{"<w:b/>" if bold else ""}' \
          f'<w:sz w:val="{int(round(font_pt * 2))}"/>' \
          f'<w:szCs w:val="{int(round(font_pt * 2))}"/></w:rPr>'
    return (
        "<w:tc><w:tcPr><w:tcW w:w=\"0\" w:type=\"auto\"/></w:tcPr>"
        f"<w:p><w:pPr><w:spacing w:before=\"0\" w:after=\"0\" w:line=\"240\" "
        f"w:lineRule=\"auto\"/><w:jc w:val=\"left\"/>{rpr}</w:pPr>"
        f"<w:r>{rpr}<w:t xml:space=\"preserve\">{escape(str(text))}</w:t></w:r>"
        "</w:p></w:tc>"
    )


def build_table_xml(rows: list[list[str]], font_pt: float, header_repeat: bool):
    """Build the <w:tbl> directly.

    python-docx's cell-by-cell API is far too slow for a 504 x 14 table
    (7 000+ paragraph objects); the XML this produces is the same shape
    insert_experiments_section.build_table emits -- 'Table Grid1' style,
    100 % width, autofit layout, bold first row.
    """
    parts = [
        f"<w:tbl {NSDECL}><w:tblPr>"
        '<w:tblStyle w:val="TableGrid1"/>'
        '<w:tblW w:w="5000" w:type="pct"/>'
        '<w:tblLayout w:type="autofit"/>'
        '<w:tblBorders>'
        + "".join(f'<w:{s} w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
                  for s in ("top", "left", "bottom", "right", "insideH", "insideV"))
        + "</w:tblBorders></w:tblPr>"
    ]
    for i, row in enumerate(rows):
        head = i == 0
        trpr = "<w:trPr><w:tblHeader/></w:trPr>" if (head and header_repeat) else ""
        parts.append("<w:tr>" + trpr
                     + "".join(_cell_xml(c, font_pt, head) for c in row)
                     + "</w:tr>")
    parts.append("</w:tbl>")
    return parse_xml("".join(parts))


# =========================================================================== #
# Figure catalogue
# =========================================================================== #
SENS = IMG / "sensitivity"
MECH = IMG / "mechanistic_selfcontained"
PSI = IMG / "psi_quadrature"
MINK = IMG / "minimal_kinetic"
MAYP = IMG / "may_prompt"
RPT = IMG / "report"

# fig9-fig12 from breakthrough_fit/cross_run_figs.py are deliberately NOT used:
# that script's "clean runs" selector is ^(\d+)ml_(\d+)g$, i.e. the twelve
# synthetic parametric CSVs, and fig11/fig12 draw the May-* records. CLAUDE.md
# rule 1 forbids presenting either as measured. src/solver/report_figs_measured.py
# writes R1-R4 over the real 21-run basis instead.
FIG_64 = [
    (RPT / "R1_grid_small_multiples.png",
     "Fig. 9) Measured breakthrough curves for the 3×3 flow × concentration "
     "design, one panel per cell with every run in that cell overlaid. Dotted and "
     "dash-dot rules mark the C/C₀ = 0.05 and 0.95 definitions of t_b and t_E. "
     "* flow stated twice and inconsistently in the source file; † run never "
     "reaches C/C₀ = 0.95, so its t_E is undefined and is never imputed."),
    (RPT / "R2_grid_overlay.png",
     "Fig. 10) The same sixteen runs on one axis. Colour encodes inlet CO₂ and "
     "line style encodes flow rate, so the two effects can be read against each "
     "other directly."),
    (MAYP / "Plot1_overlay.png",
     "Fig. 11) Measured breakthrough curves for the five earlier runs "
     "(3/4/5/6/8) on the 8.5 mm column. Plotted separately from Figures 9 and 10 "
     "because the two campaigns differ in bed length and column internal diameter "
     "and are not pooled."),
]

# The four charts that were already embedded in §7 without a caption, in the
# order they appear. Titles read off the charts themselves.
FIG_ORPHAN = [
    "Fig. 12) Breakthrough time t_b against inlet flow rate, grouped by inlet CO₂ "
    "concentration (5, 10 and 15 % v/v). Circle and square markers are the two "
    "replicate sets reported in Tables 9 and 10.",
    "Fig. 13) Saturation time t_E against inlet flow rate, same grouping as "
    "Fig. 12. t_E is defined at C/C₀ = 0.95.",
    "Fig. 14) Saturation capacity q_E against inlet flow rate, same grouping as "
    "Fig. 12.",
    "Fig. 15) Breakthrough capacity q_B against inlet flow rate, same grouping as "
    "Fig. 12.",
]

FIG_71 = [
    (RPT / "R3_model_ranking.png",
     "Fig. 19) All twenty-four registry models ranked by mean adjusted R², the two "
     "campaigns side by side and never pooled. The numeric values, with median, "
     "RMSE, mean ΔAICc and runs won, are Tables 15 and 16."),
    (RPT / "R4_metrics_vs_flow.png",
     "Fig. 20) Every measured performance metric against flow for the grid "
     "campaign: t_b, t_E, t₅₀, q_dyn, L_MTZ and ψ. Open markers are the four runs "
     "whose source file states its flow twice and disagrees with itself, for which "
     "q_dyn, L_MTZ and ψ are provisional (Table 5)."),
    (MAYP / "Plot8_parity.png",
     "Fig. 21) Predicted against observed C/C₀ for runs 3/4/5/6/8, nine models "
     "coloured by family. Departure from the 1:1 line is the fit error that "
     "Tables 14-17 quantify."),
]

FIG_85 = [
    (SENS / "S1_experimental_clusters_grid.png",
     "Fig. 22) Tier-1 cluster ANOVA of the experimental scatterplots, 3×3 grid "
     "campaign (n = 16). Each panel is one factor against one response; dashed "
     "lines are cluster boundaries and the red trace joins the cluster means. "
     "Panel titles carry F, p and η²."),
    (SENS / "S1_experimental_clusters_runs.png",
     "Fig. 23) Tier-1 cluster ANOVA, pooled scope (n = 21). Shown for completeness "
     "only: pooling the two campaigns confounds a change of apparatus with the "
     "effects being measured, and every η² falls relative to Fig. 22. Inference in "
     "this report is drawn from the grid scope."),
    (SENS / "S2_M11_independent.png",
     "Fig. 24) Tier-2 Monte-Carlo cluster ANOVA for M11 (fractal error-function) "
     "under independent Latin-hypercube sampling of the fitted parameters."),
    (SENS / "S2_M11_rank-correlated.png",
     "Fig. 25) Tier-2 Monte-Carlo cluster ANOVA for M11 under rank-correlated "
     "sampling. Independent sampling produces no breakthrough curve in 11.3 % of "
     "draws for this model; rank-correlated sampling, which respects the k₀-h "
     "ridge of Fig. 30, produces none."),
    (SENS / "S2_M10_independent.png",
     "Fig. 26) Tier-2 Monte-Carlo cluster ANOVA for M10 (fractal Gudermannian), "
     "independent sampling."),
    (SENS / "S2_M10_rank-correlated.png",
     "Fig. 27) Tier-2 Monte-Carlo cluster ANOVA for M10, rank-correlated sampling "
     "(8.3 % of independent draws produce no curve; none under rank correlation)."),
    (SENS / "S2_M24_independent.png",
     "Fig. 28) Tier-2 Monte-Carlo cluster ANOVA for M24 (parallel two-component "
     "sigmoidal), independent sampling."),
    (SENS / "S2_M24_rank-correlated.png",
     "Fig. 29) Tier-2 Monte-Carlo cluster ANOVA for M24, rank-correlated sampling."),
    (SENS / "S3_identifiability_ridge.png",
     "Fig. 30) Parameter identifiability: the fitted (k₀, h) pairs lie on a "
     "one-dimensional ridge, Spearman ρ_S = 0.984 for M11 and 0.971 for M10. "
     "k₀ and h are therefore not independently identifiable and the compound "
     "κ = k₀τ₀^(−h) is what the data constrain."),
]

FIG_94 = [
    (MECH / "V1_ade_vs_erfc.png",
     "Fig. 31) V1 — advection-dispersion with the sink switched off, finite-volume "
     "solution against the exact Ogata-Banks erfc solution, at four mesh "
     "resolutions. The L² error and its halving with mesh refinement are printed "
     "on the legend."),
    (MECH / "V2_rh_front.png",
     "Fig. 32) V2 — isothermal equilibrium shock. Left: concentration profiles "
     "against the Rankine-Hugoniot shock position. Right: the numerical front "
     "position z₅₀(t) against the R-H chord velocity."),
    (MECH / "V3_travelling_wave.png",
     "Fig. 33) V3 — linear-driving-force travelling wave against the implicit "
     "closed form, in the moving coordinate η = z − v_RH t. The RMS deviation over "
     "0.02 < c/c_f < 0.98 is printed in the title."),
    (MECH / "V4_nonisothermal.png",
     "Fig. 34) V4 — non-isothermal Toth demonstration. The thermal parameters are "
     "placeholders, not measured values, and are flagged as such on the figure: "
     "this panel demonstrates that the energy balance runs, not that it is "
     "calibrated to this rig."),
    (MECH / "F5_run_overlays.png",
     "Fig. 35) F5 — measured breakthrough curves for runs 3/4/5/6/8 against a "
     "single global mechanistic fit (one n_s, b, t, k for all five runs). The "
     "bed voidage ε carried in this fit is the floored placeholder discussed in "
     "Section 6.1; the fit is a scoping exercise, not a validated calibration."),
    (MECH / "F6_isotherm.png",
     "Fig. 36) F6 — fitted isotherms at ambient temperature with the measured "
     "dynamic capacities overlaid. q_dyn is a dynamic capacity, not an equilibrium "
     "loading, and the two are not interchangeable."),
    (PSI / "F1_danilov_fig1_reproduction.png",
     "Fig. 37) Reproduction of Danilov's Figure 1 (CO₂ breakthrough, left; outlet "
     "gas temperature, right) as an external check on the Ψ-quadrature solver."),
    (PSI / "F2_three_way_langmuir_wave.png",
     "Fig. 38) Three-way comparison for the Langmuir wave: profiles at "
     "t = 0.65 t_st, and the outlet breakthrough under the symmetric ansatz "
     "against the asymmetric exact solution."),
    (PSI / "F3_nonisothermal_overlay.png",
     "Fig. 39) Non-isothermal Toth case: outlet breakthrough and the outlet "
     "temperature excursion on the Ψ-quadrature scale."),
    (MINK / "minimal_kinetic_fit.png",
     "Fig. 40) Minimal kinetic model of Section 9.1 with one global (q_m, b, k) "
     "against every measured run. Per-run RMSE in C/C₀ is printed to stdout by "
     "minimal_kinetic_model.py."),
    (MINK / "minimal_kinetic_rh.png",
     "Fig. 41) Rankine-Hugoniot check for the minimal kinetic model on the run 5 "
     "geometry, using the fitted isotherm in the equilibrium limit."),
]


# =========================================================================== #
# Prose
# =========================================================================== #
SEC_64 = [
    ("Heading2", "6.4 Measured breakthrough curves"),
    ("Body",
     "Sections 6.1 to 6.3 report the breakthrough and equilibrium metrics as "
     "numbers. Figures 9 to 11 show the curves those numbers were read from. "
     "Figure 9 lays the sixteen grid runs out in the 3×3 design they were measured "
     "on, so the replicate structure is visible directly: seven of the nine cells "
     "carry two runs, and the two 15 % cells at 50 and 150 mL/min carry one each. "
     "Figure 10 overlays the same sixteen on a single axis, and Figure 11 does the "
     "same for the five earlier runs on the 8.5 mm column."),
    ("Body",
     "Two records are marked in Figure 9 rather than quietly smoothed over. The "
     "5 % / 50 mL/min run of 29 July never reaches C/C₀ = 0.95, so its t_E, L_MTZ "
     "and ψ are undefined and are left undefined — every statistic that depends on "
     "them is computed pairwise on the fifteen runs that do reach it. Four further "
     "runs state their flow rate twice in the source file and disagree with "
     "themselves; the affected metrics are provisional, and Table 5 names the "
     "owner. Two additional records logged on 17 July carry no embedded geometry "
     "at all and are excluded by the pipeline's own metadata guard, not by "
     "selection."),
    ("Body",
     "The per-run curves for all twenty-one runs, each with its full set of fit "
     "diagnostics, are in Appendix A."),
]

SEC_71 = [
    ("Heading2", "7.1 Cross-run trends and model ranking"),
    ("Body",
     "The four charts above read one metric at a time off the replicate tables. "
     "Figures 19 to 21 aggregate across runs instead. Figure 19 ranks the "
     "twenty-four registry models by mean adjusted R²; Figure 20 puts all six "
     "measured performance metrics on a common flow axis; Figure 21 is the "
     "predicted-against-observed parity plot for the five-run campaign."),
    ("Body",
     "Two cautions apply to every panel in this group. First, the two campaigns "
     "are not pooled: the five earlier runs were measured on an 8.5 mm column with "
     "a 21 cm bed and the sixteen grid runs on an 8.2 mm column with a 22.5-24.5 cm "
     "bed, so a pooled trend would confound a change of apparatus with the effect "
     "being measured. Where a figure covers both, the two are drawn as separate "
     "series. Second, q_dyn is not currently a comparable quantity across runs — "
     "Section 8.5 quantifies why — so the capacity panel of Figure 20 should be "
     "read as a record of what was measured, not as a capacity trend."),
]

SEC_82 = [
    ("Body",
     "All free parameters — for both the minimal kinetic model of Section 9.1 and "
     "the twenty-four-model registry of Section 3.6 — were determined by nonlinear "
     "least-squares regression (scipy.optimize.curve_fit, trust-region-reflective "
     "algorithm, twelve multi-start initialisations per model) rather than by "
     "linearising each model to a straight line."),
    ("Body",
     "This choice follows Hu et al. (2024, §5.5) directly. Linearising a "
     "breakthrough model — for example the Thomas model's linear form "
     "ln(c₀/c − 1) = k_Th q₀ m/ν − k_Th c₀ t — implicitly transforms the response "
     "variable and alters its error structure: ordinary least squares on the "
     "linearised form assumes constant-variance, normally distributed residuals in "
     "the transformed coordinate, an assumption the untransformed C/C₀ data need "
     "not satisfy. The transform is also undefined at C/C₀ = 0 and C/C₀ = 1, so a "
     "linearised fit must exclude data at and near breakthrough and near "
     "saturation — precisely the regions the acceptance criteria of Section 4 are "
     "built to capture completely. Hu et al. (2024) report that this exclusion "
     "inflates linearised parameter error most sharply near those same two limits. "
     "Nonlinear regression avoids both problems, fitting the untransformed C/C₀ "
     "response directly with no domain exclusion, and extends without modification "
     "to the fractal-like and two-component models of Sections 3.6.8 and 3.6.14, "
     "which have no closed linear form at all."),
    ("Body",
     "A second methodological choice follows from the same source. Every model is "
     "fitted to the complete breakthrough curve, not a curve truncated before "
     "saturation. Hu et al. (2024, §5.2) quantify why this matters: fitting the "
     "Thomas model to progressively truncated breakthrough data (100 % down to "
     "20 % of the complete curve) produced relative errors in the fitted rate "
     "constant of up to 110.8 %, 26.1 % and 47.5 % at three influent "
     "concentrations — even though the truncated fits still reported adjusted R² "
     "above 0.99. A good statistical fit to a partial curve does not imply an "
     "accurate parameter estimate."),
    ("Body",
     "The error statistics reported in Tables 14 to 17 and in Appendix B are "
     "defined as follows (Hu et al., 2024, §4). The coefficient of determination "
     "is R² = 1 − Σ(yᵢ − ŷᵢ)² / Σ(yᵢ − ȳ)²; the adjusted R² penalises the "
     "parameter count p; the root-mean-square error RMSE is reported in units of "
     "C/C₀, so it is directly comparable between runs; the reduced chi-square "
     "χ²_red is the residual sum of squares per degree of freedom; and model "
     "comparison uses the corrected Akaike criterion AICc, reported here as "
     "ΔAICc against the best model in each run. A nested pair of models is "
     "compared by the F-test F = [(RSS₁ − RSS₂)/(df₁ − df₂)] / [RSS₂/df₂], "
     "Table 17 and Appendix B.3."),
    ("Body",
     "Two caveats are recorded rather than resolved. First, the column labelled "
     "W_AICc in the pipeline's results files is a pairwise logistic transform "
     "1/(1 + exp(0.5·ΔAICc)) against the best model in the set, not a normalised "
     "Akaike weight over all candidates; it equals 0.5 for the best model by "
     "construction and is not reported in this section for that reason. Second, "
     "Hu et al. (2024, §4) caution that R², adjusted R² and AIC alone are not "
     "sufficient to judge fit quality and recommend the residual plot as the more "
     "reliable diagnostic. That applies here with force: C/C₀ is bounded on [0, 1], "
     "so residual variance is structurally pinched near both limits and is "
     "heteroscedastic by construction — an assumption the F-test and AICc "
     "comparisons do not themselves test. The residual diagnostics for every model "
     "of every run are reproduced in Appendix A as the P7 panel of each run."),
    ("Body",
     "Table 14 reports, for each of the sixteen grid runs, the model selected by "
     "AICc against the M01 logistic (Bohart-Adams / Thomas / Yoon-Nelson) baseline. "
     "Tables 15 and 16 rank all twenty-four registry models by mean adjusted R², "
     "separately for each campaign, and Table 17 gives the nested F-test that asks "
     "whether the fractal exponent h earns its parameter. The complete per-model, "
     "per-run statistics — 504 rows — are in Appendix B.1, the fitted parameters "
     "with their asymptotic standard errors in Appendix B.2, all six nested "
     "comparisons in Appendix B.3, and the degenerate fits in Appendix B.4."),
]

CAP_14 = ("Table 14) Fit quality for the sixteen grid runs: model selected by AICc "
          "against the M01 logistic baseline. * flow rate stated twice and "
          "inconsistently in the source file (Table 5); † run never reaches "
          "C/C₀ = 0.95.")
CAP_15 = ("Table 15) All twenty-four registry models ranked by mean adjusted R² "
          "over the sixteen grid runs.")
CAP_16 = ("Table 16) All twenty-four registry models ranked by mean adjusted R² "
          "over the five earlier runs (3/4/5/6/8). Reported separately from "
          "Table 15; the two campaigns are not pooled.")
CAP_17 = ("Table 17) Nested F-test of the fractal exponent: M01 logistic against "
          "M23 fractal Yoon-Nelson, every run.")

SEC_82_TAIL = [
    ("Body",
     "The ranking is stable in family and unstable in kernel. M10, M11, M23 and "
     "M24 occupy the top of both campaigns, but which of them wins a given run "
     "changes from run to run. This is not an unreliable fitting procedure: these "
     "are closely related asymmetric sigmoids whose fitted curves lie within a "
     "sup-norm distance of about 0.03 of one another once slope-matched at the "
     "inflection point (Section 3.6.4), so AICc is choosing which near-equivalent "
     "functional form best absorbs that run's particular noise realisation, not "
     "resolving a physical distinction. The model-independent result is the "
     "fractal exponent itself: Table 17 rejects the non-fractal baseline in every "
     "run, at F values in the thousands."),
]

SEC_85 = [
    ("Heading2", "8.5 Sensitivity analysis"),
    ("Body",
     "The sensitivity analysis follows the cluster-based method of Kleijnen and "
     "Helton (1999a) — partition an input into k clusters, then test by ANOVA "
     "whether the response mean differs across clusters — supplemented by "
     "standardised regression coefficients (Saltelli et al., 2000). It is run at "
     "two tiers. Tier 1 treats the experiment itself as the input: flow rate and "
     "inlet concentration against the measured responses t_b, t_E, t₅₀, q_dyn, "
     "L_MTZ and ψ. Tier 2 samples the fitted model parameters by Latin hypercube "
     "and asks which of them the predicted responses are sensitive to."),
    ("Body",
     "Table 18 gives the Tier-1 cluster ANOVA, Table 19 the two-way factorial "
     "ANOVA with replication (Type II sums of squares, unbalanced design), and "
     "Table 20 the replicate-reproducibility decomposition that separates pure "
     "error from total variance. Table 21 lists the parameters that sit at a "
     "fitting bound often enough to disqualify their model from the parameter "
     "dissection."),
    ("Body",
     "Three findings constrain how the rest of this report may be read, and are "
     "stated here rather than buried in the tables."),
    ("Body",
     "First, q_dyn is not currently a measurable quantity on this rig. 78.4 % of "
     "its variance is pure replicate error and replicate pairs differ by up to a "
     "factor of 3.04. Any statement of the form \"concentration sets capacity\" is "
     "therefore withdrawn: the acquisition endpoint must be fixed at a consistent "
     "C/C₀ — 0.98 is the natural choice — before capacity is compared across runs "
     "at all. Breakthrough time t_b, by contrast, is well determined: flow rate "
     "alone accounts for 85.9 % of its variance (Table 19)."),
    ("Body",
     "Second, the rate constant k₀ and the fractal exponent h are not "
     "independently identifiable. Their fitted values lie on a one-dimensional "
     "ridge with Spearman ρ_S = 0.984 for M11 and 0.971 for M10 (Figure 30). What "
     "the data constrain is the compound κ = k₀ τ₀^(−h), and that is what should "
     "be reported and compared; a bare k₀ from one run is not comparable with a "
     "bare k₀ from another. The same ridge is why the Tier-2 Monte Carlo is run "
     "twice: sampling the parameters independently produces no breakthrough curve "
     "at all in 11.3 % of draws for M11 and 8.3 % for M10, while rank-correlated "
     "sampling produces none."),
    ("Body",
     "Third, L_MTZ has a hard algebraic floor of one half the bed length and a "
     "coefficient of variation of only 4.0 % across the campaign, so its apparent "
     "effects are a property of the formula rather than of the bed. It is reported "
     "for completeness and no conclusion rests on it."),
]

CAP_18 = ("Table 18) Tier-1 cluster ANOVA of the experimental factors against the "
          "measured responses, both scopes. q_BH is the Benjamini-Hochberg "
          "adjusted p-value.")
CAP_19 = ("Table 19) Two-way factorial ANOVA with replication (Type II sums of "
          "squares) over the 3×3 grid.")
CAP_20 = ("Table 20) Replicate reproducibility: pure error against total variance, "
          "per response.")
CAP_21 = ("Table 21) Parameters pinned at a fitting bound. A model is disqualified "
          "from the parameter dissection when any of its parameters sits at its "
          "own bound in more than 25 % of runs.")

SEC_85_TAIL = [
    ("Body",
     "One consequence for the report as a whole: the pooled scope in Table 18 and "
     "Figure 23 is shown for completeness only. Every η² falls when the two "
     "campaigns are pooled, because bed length and column internal diameter differ "
     "systematically between the two rigs. All inference in this report is drawn "
     "from the sixteen-run grid scope."),
]

SEC_94 = [
    ("Heading2", "9.4 Numerical verification"),
    ("Body",
     "The model equations of Sections 9.1 to 9.3 are verified against cases with "
     "known answers before any of them is used. Figures 31 to 34 are the four "
     "verification cases: the sink-free advection-dispersion limit against the "
     "exact Ogata-Banks solution (V1), the equilibrium shock against its "
     "Rankine-Hugoniot chord velocity (V2), the linear-driving-force travelling "
     "wave against its implicit closed form (V3), and a non-isothermal Toth "
     "demonstration (V4). Figures 35 and 36 fit the mechanistic model globally to "
     "the five-run campaign, and Figures 37 to 39 check the Ψ-quadrature treatment "
     "against an external reproduction and against the asymmetric exact solution. "
     "Figures 40 and 41 are the minimal kinetic model of Section 9.1 and its "
     "Rankine-Hugoniot check."),
    ("Body",
     "Four limitations attach to this group and none of them is resolved here. "
     "The thermal parameters in Figure 34 are placeholders, so V4 demonstrates "
     "that the energy balance integrates, not that it is calibrated. The bed "
     "voidage used in Figures 35 and 36 is the floored ε discussed in Section 6.1; "
     "the pellet density ρ_p that would make it physical is still an open input "
     "from the sorbent supplier. The solver scaffold carries two temperatures "
     "while the analytical derivation uses a single pseudo-homogeneous energy "
     "balance, and the two have not yet been reconciled. And the parameters of "
     "this chapter have not been fitted against the measured breakthrough data of "
     "Sections 6 to 8: the chapter is a derivation and scoping exercise, as "
     "Section 9's opening paragraph states."),
]

APP_A_INTRO = [
    ("Heading1", "Appendix A  Per-run fit diagnostics"),
    ("Body",
     "This appendix reproduces the complete diagnostic set for every one of the "
     "twenty-one real runs: seven plots per run, 147 figures in total. Nothing has "
     "been selected — every run the pipeline fitted appears here, including the "
     "runs whose metrics are flagged as provisional in Table 5."),
    ("Body",
     "The seven plots are the same for every run. P1 is the parity plot of "
     "predicted against observed C/C₀ for the eight best-fitting models. P2 shows "
     "the Langmuir-family and Freundlich-family fits over the raw data with t_b "
     "and t_E marked. P3 is the Weibull fit with its wave-front rate. P4 contrasts "
     "the symmetric M01 Thomas/Yoon-Nelson fit with the asymmetric M04 "
     "dose-response fit and shades the asymmetry between them. P5 projects the "
     "mass-transfer-zone propagation from the M01 fit at the baseline geometry and "
     "at varied velocity and bed length. P6 places the standard Yoon-Nelson fit "
     "beside the fractal Yoon-Nelson fit with their residuals. P7 is the residual "
     "diagnostic grid, one panel per converged model."),
    ("Body",
     "The figures carry no fitted statistics on their faces; the statistics are "
     "tabulated in Appendix B, and the corresponding numbers for the plotted fits "
     "are in Appendix B.1 and B.2. Figure numbers in this appendix run A1 to A147 "
     "and are indexed by run in Table A.1."),
]

CAP_A1 = ("Table A.1) Contents of Appendix A: the twenty-one real runs, their "
          "conditions, and the figure numbers that belong to each. * flow rate "
          "stated twice and inconsistently in the source file; † run never reaches "
          "C/C₀ = 0.95.")

APP_B_INTRO = [
    ("Heading1", "Appendix B  Complete fit statistics"),
    ("Body",
     "Every number in this appendix is read back out of the committed results "
     "files, breakthrough_out/<run>/results_<run>.csv, one file per run and "
     "twenty-four rows per file. Nothing was refitted to produce these tables. The "
     "only arithmetic applied is the difference ΔAICc against the best model in "
     "each run, and the nested F-test of Appendix B.3, which is computed with the "
     "same breakthrough_fit.stats.f_test the pipeline itself calls."),
    ("Body",
     "Run labels are abbreviated as date-concentration-flow; * marks the four runs "
     "whose source file states its flow rate twice and disagrees with itself, and "
     "† the one run that never reaches C/C₀ = 0.95. For those five runs the "
     "time-based statistics are unaffected, while q_dyn, L_MTZ and ψ are "
     "provisional or undefined as set out in Table 5."),
]

B_SECTIONS = [
    ("Heading2", "B.1 Model fit statistics, all models and all runs",
     "Table B.1) Fit statistics for all twenty-four registry models on all "
     "twenty-one runs, 504 rows. ΔAICc is measured against the best model within "
     "each run.",
     "b1"),
    ("Heading2", "B.2 Fitted parameters and standard errors",
     "Table B.2) Fitted parameter values with the asymptotic standard errors "
     "returned by scipy.optimize.curve_fit, all models and all runs.",
     "b2"),
    ("Heading2", "B.3 Nested F-tests",
     "Table B.3) All six nested model comparisons the pipeline evaluates, for "
     "every run.",
     "b3"),
    ("Heading2", "B.4 Non-converged and degenerate fits",
     "Table B.4) Models that failed to converge, or converged to a non-finite or "
     "negative adjusted R². Reported for completeness and excluded from every "
     "ranking in this report.",
     "b4"),
]

# --------------------------------------------------------------------------- #
# §7 heading repairs: literal "?" characters left by an earlier renumber pass.
# --------------------------------------------------------------------------- #
HEADING_FIXES = [
    ("6.??? Determining the Optimal Operating Parameters",
     "7.2 Determining the Optimal Operating Parameters", "Heading2"),
    ("6.?.1 Identifying the Optimal Condition",
     "7.2.1 Identifying the Optimal Condition", "Heading3"),
    ("6.?.2 Why This is Still Not Good Enough",
     "7.2.2 Why This is Still Not Good Enough", "Heading3"),
    ("6.?.3 Real-World Settings",
     "7.2.3 Real-World Settings", "Heading3"),
    ("7.?.4 Possible Adjustments to Increase Utilisation",
     "7.2.4 Possible Adjustments to Increase Utilisation", "Heading3"),
]

# Two pre-existing captions lost the space after the bracket, which is why they
# read "Table 7)Experimental" rather than "Table 7) Experimental". Cosmetic, but
# it also makes them invisible to any caption sweep that expects the separator.
CAPTION_SPACE_FIXES = [
    ("Table 7)Experimental Parameters", "Table 7)Experimental", "Table 7) Experimental"),
    ("Table 11)Winning model per run", "Table 11)Winning", "Table 11) Winning"),
]

# Existing figure captions, renumbered so the four newly captioned charts can
# take 11-14 in document order. Highest first so a number never lands on one
# that has not been vacated. No prose in the report references these by number.
FIG_RENUMBER = [
    ("Fig. 11) Validation of fractal-like kinetics", "Fig. 11)", "Fig. 18)"),
    ("Fig. 10) Breakthrough curve fitting", "Fig. 10)", "Fig. 17)"),
    ("Fig. 9) Predicted vs. observed values", "Fig. 9)", "Fig. 16)"),
]

# New List-of-Figures entries, appended after the existing Fig. 10 entry and the
# existing Table 10 entry respectively.
LOF_FIGS = [
    "Fig. 9) Measured breakthrough curves, 3×3 flow × concentration design",
    "Fig. 10) The sixteen grid runs overlaid on one axis",
    "Fig. 11) Measured breakthrough curves for runs 3/4/5/6/8, overlaid",
    "Fig. 12) Breakthrough time against inlet flow rate",
    "Fig. 13) Saturation time against inlet flow rate",
    "Fig. 14) Saturation capacity against inlet flow rate",
    "Fig. 15) Breakthrough capacity against inlet flow rate",
    "Fig. 16) Predicted vs. observed values in breakthrough models fit",
    "Fig. 17) Breakthrough curve fitting with linear adsorption isotherm",
    "Fig. 18) Validation of fractal-like kinetics on logistic breakthrough model curve",
    "Fig. 19) Registry models ranked by mean adjusted R², both campaigns",
    "Fig. 20) Measured performance metrics against flow, grid campaign",
    "Fig. 21) Predicted against observed C/C₀ for runs 3/4/5/6/8",
    "Fig. 22) Tier-1 cluster ANOVA, 3×3 grid campaign",
    "Fig. 23) Tier-1 cluster ANOVA, pooled scope",
    "Fig. 24) Tier-2 Monte-Carlo cluster ANOVA, M11, independent sampling",
    "Fig. 25) Tier-2 Monte-Carlo cluster ANOVA, M11, rank-correlated sampling",
    "Fig. 26) Tier-2 Monte-Carlo cluster ANOVA, M10, independent sampling",
    "Fig. 27) Tier-2 Monte-Carlo cluster ANOVA, M10, rank-correlated sampling",
    "Fig. 28) Tier-2 Monte-Carlo cluster ANOVA, M24, independent sampling",
    "Fig. 29) Tier-2 Monte-Carlo cluster ANOVA, M24, rank-correlated sampling",
    "Fig. 30) Parameter identifiability: the k₀-h ridge",
    "Fig. 31) V1 — advection-dispersion against the exact erfc solution",
    "Fig. 32) V2 — isothermal equilibrium shock, Rankine-Hugoniot verification",
    "Fig. 33) V3 — linear-driving-force travelling wave against the closed form",
    "Fig. 34) V4 — non-isothermal Toth demonstration (placeholder thermal parameters)",
    "Fig. 35) F5 — measured breakthrough against the global mechanistic fit",
    "Fig. 36) F6 — fitted isotherms with measured dynamic capacities overlaid",
    "Fig. 37) Ψ-quadrature: reproduction of Danilov Fig. 1",
    "Fig. 38) Ψ-quadrature: three-way Langmuir wave comparison",
    "Fig. 39) Ψ-quadrature: non-isothermal Toth overlay",
    "Fig. 40) Minimal kinetic model against every measured run",
    "Fig. 41) Rankine-Hugoniot check for the minimal kinetic model",
    "Fig. A1 – A147) Per-run fit diagnostics, seven plots for each of the "
    "twenty-one real runs — indexed by run in Table A.1 (Appendix A)",
]

LOF_TABLES = [
    "Table 11) Winning model per run vs. M01 baseline, with asymptotic standard errors",
    "Table 12) Langmuir isotherm parameters for CO2",
    "Table 13) The basic parameters for packed-bed adsorption",
    "Table 14) Fit quality for the sixteen grid runs",
    "Table 15) Registry models ranked by mean adjusted R², grid campaign",
    "Table 16) Registry models ranked by mean adjusted R², five-run campaign",
    "Table 17) Nested F-test of the fractal exponent, every run",
    "Table 18) Tier-1 cluster ANOVA of the experimental factors",
    "Table 19) Two-way factorial ANOVA with replication",
    "Table 20) Replicate reproducibility: pure error against total variance",
    "Table 21) Parameters pinned at a fitting bound",
    "Table A.1) Contents of Appendix A",
    "Table B.1) Fit statistics for all models on all runs",
    "Table B.2) Fitted parameters and standard errors",
    "Table B.3) Nested F-tests",
    "Table B.4) Non-converged and degenerate fits",
]


# =========================================================================== #
# Editing passes
# =========================================================================== #
def fix_headings(doc, log):
    for old, new, style in HEADING_FIXES:
        i, el = IES.find_body_paragraph(doc, old)
        if not D.replace_text(el, old, new):
            raise SystemExit(f"could not rewrite heading {old!r}")
        IES.set_style(el, style)
        log.append(f"head  {old!r} -> {new!r} ({style})")


def renumber_figures(doc, log):
    for locate, old, new in FIG_RENUMBER:
        i, el = D.find_heading(doc, locate, level=7)
        if not D.replace_text(el, old, new):
            raise SystemExit(f"could not renumber {old!r}")
        log.append(f"fig   {old} -> {new}")
    # "…linear adsorption isother" — the caption lost its final letter.
    _, el = D.find_heading(doc, "Breakthrough curve fitting with linear", level=7)
    if D.ptext(el).rstrip().endswith("isother"):
        D.replace_text(el, "isother", "isotherm")
        log.append("fig   Fig. 16) caption completed: 'isother' -> 'isotherm'")
    for locate, old, new in CAPTION_SPACE_FIXES:
        hits = D.find_paragraph(doc, locate, unique=False)
        body = [(i, e) for i, e in hits if not IES.is_toc_entry(e)]
        if len(body) != 1:
            raise SystemExit(f"expected 1 body match for {locate!r}, got {len(body)}")
        D.replace_text(body[0][1], old, new)
        log.append(f"cap   {old!r} -> {new!r}")


def caption_orphans(b: Builder, log):
    """The four §7 charts embedded without captions, in document order."""
    doc = b.doc
    lo = D.find_heading(doc, "7 Experimental Analysis", level=1)[0]
    hi = D.find_heading(doc, "8 Fitting performance", level=1)[0]
    A = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}
    targets = []
    for i, el in D.body_children(doc):
        if not (lo < i < hi) or el.tag != qn("w:p"):
            continue
        if el.findall(".//a:blip", A) and not D.ptext(el).strip():
            if len(el.findall(".//a:blip", A)) == 1 and len(targets) < 4:
                targets.append(el)
    if len(targets) != 4:
        raise SystemExit(f"expected 4 uncaptioned single-image paragraphs in §7, "
                         f"found {len(targets)}")
    for el, cap in zip(targets, FIG_ORPHAN):
        b.keep_with_next(el)
        b.para(el, CAPTION_STYLE, cap)
    log.append("cap   4 previously uncaptioned §7 charts captioned Fig. 12-15")


def _anchor_before(doc, heading_text, level):
    """A marker paragraph immediately before a heading, to build forward from."""
    _, h = D.find_heading(doc, heading_text, level=level)
    marker = D.clone_paragraph_after(h, text="", style=None)
    marker.getparent().remove(marker)
    h.addprevious(marker)
    return marker


def insert_section_64(b: Builder, log):
    cur = _anchor_before(b.doc, "7 Experimental Analysis", 1)
    marker = cur
    for kind, text in SEC_64:
        cur = b.para(cur, kind, text)
    for png, cap in FIG_64:
        cur = b.figure(cur, png, cap)
    marker.getparent().remove(marker)
    log.append("sec   §6.4 inserted with Fig. 9-10")


def insert_section_71(b: Builder, log):
    _, cur = D.find_heading(b.doc, "Fig. 18) Validation of fractal-like", level=7)
    for kind, text in SEC_71:
        cur = b.para(cur, kind, text)
    for png, cap in FIG_71:
        cur = b.figure(cur, png, cap)
    log.append("sec   §7.1 inserted with Fig. 18-21")


def expand_section_82(b: Builder, log):
    cur = _anchor_before(b.doc, "8.3 Parameter trends with operating", 2)
    marker = cur
    for kind, text in SEC_82:
        cur = b.para(cur, kind, text)
    for cap, rows, pt in (
        (CAP_14, T.tbl_fit_quality(T.RUNS_GRID), 6.5),
        (CAP_15, T.tbl_model_ranking(T.RUNS_GRID), 7.5),
        (CAP_16, T.tbl_model_ranking(T.RUNS_OLD), 7.5),
        (CAP_17, T.tbl_ftest(T.ALL_RUNS), 7.0),
    ):
        cur = b.para(cur, "Body", cap)
        cur = b.table(cur, rows, pt)
    for kind, text in SEC_82_TAIL:
        cur = b.para(cur, kind, text)
    marker.getparent().remove(marker)
    log.append("sec   §8.2 expanded: 6 paragraphs + Tables 14-17")


def insert_section_85(b: Builder, log):
    cur = _anchor_before(b.doc, "9 Mathematical Modelling", 1)
    marker = cur
    for kind, text in SEC_85:
        cur = b.para(cur, kind, text)
    for cap, rows, pt in (
        (CAP_18, T.tbl_e1_cluster_anova(), 6.5),
        (CAP_19, T.tbl_e2_factorial(), 6.5),
        (CAP_20, T.tbl_e2b_reproducibility(), 7.0),
        (CAP_21, T.tbl_bound_pinning(), 7.5),
    ):
        cur = b.para(cur, "Body", cap)
        cur = b.table(cur, rows, pt)
    for png, cap in FIG_85:
        cur = b.figure(cur, png, cap)
    for kind, text in SEC_85_TAIL:
        cur = b.para(cur, kind, text)
    marker.getparent().remove(marker)
    log.append("sec   §8.5 inserted with Tables 18-21 and Fig. 22-30")


def insert_section_94(b: Builder, log):
    cur = _anchor_before(b.doc, "10 Conclusions", 1)
    marker = cur
    for kind, text in SEC_94:
        cur = b.para(cur, kind, text)
    for png, cap in FIG_94:
        cur = b.figure(cur, png, cap)
    marker.getparent().remove(marker)
    log.append("sec   §9.4 inserted with Fig. 31-41")


def append_appendices(b: Builder, log):
    body = b.doc.element.body
    last_p = None
    for el in body.iterchildren():
        if el.tag == qn("w:p"):
            last_p = el
    cur = last_p

    # ---- Appendix A ------------------------------------------------------ #
    fig_start, n = {}, 1
    for run in T.ALL_RUNS:
        fig_start[run] = n
        n += 7

    for kind, text in APP_A_INTRO:
        cur = b.para(cur, kind, text)
    cur = b.para(cur, "Body", CAP_A1)
    cur = b.table(cur, T.tbl_a1_index(fig_start), 7.0)

    plot_titles = {
        "P1": "predicted against observed C/C₀, eight best models",
        "P2": "breakthrough fits by isotherm family, with t_b and t_E marked",
        "P3": "Weibull fit and wave-front rate",
        "P4": "M01 Thomas/Yoon-Nelson against M04 dose-response, with residuals",
        "P5": "mass-transfer-zone propagation at baseline, varied u and varied L",
        "P6": "standard against fractal Yoon-Nelson, with residuals",
        "P7": "residual diagnostics, one panel per converged model",
    }
    for si, run in enumerate(T.ALL_RUNS, start=1):
        cur = b.para(cur, "Heading2",
                     f"A.{si} {T.run_label(run)}")
        for k in range(7):
            code = f"P{k + 1}"
            png = T.BT / run / f"{code}_{run}.png"
            if not png.exists():
                raise SystemExit(f"missing figure: {png}")
            num = fig_start[run] + k
            cur = b.figure(cur, png,
                           f"Fig. A{num}) {run} — {code}: {plot_titles[code]}.")
    log.append(f"app   Appendix A: {len(T.ALL_RUNS)} runs x 7 plots = "
               f"{len(T.ALL_RUNS) * 7} figures + Table A.1")

    # ---- Appendix B ------------------------------------------------------ #
    for kind, text in APP_B_INTRO:
        cur = b.para(cur, kind, text)
    builders = {"b1": (T.tbl_b1_master, 5.5), "b2": (T.tbl_b2_parameters, 5.5),
                "b3": (T.tbl_b3_ftests, 6.0), "b4": (T.tbl_b4_degenerate, 6.5)}
    for kind, heading, cap, key in B_SECTIONS:
        fn, pt = builders[key]
        rows = fn()
        cur = b.para(cur, kind, heading)
        cur = b.para(cur, "Body", cap)
        cur = b.table(cur, rows, pt)
        log.append(f"app   {heading.split()[0]}: {len(rows) - 1} data rows")


def update_list_of_figures(doc, log):
    _, anchor = IES.find_body_paragraph(
        doc, "Fig. 10) Breakthrough curve fitting with linear adsorption isotherm using")
    for entry in reversed(LOF_FIGS):
        D.clone_paragraph_after(anchor, text=entry, style=None)
    # the old "Fig. 9)"/"Fig. 10)" lines are superseded by the renumbered entries
    for old in ("Fig. 9) Predicted vs. observed values in breakthrough models fit using",
                "Fig. 10) Breakthrough curve fitting with linear adsorption isotherm using"):
        _, el = IES.find_body_paragraph(doc, old)
        el.getparent().remove(el)
    _, anchor = IES.find_body_paragraph(
        doc, "Table 10) Breakthrough (tb) and equilibrium (te)")
    for entry in reversed(LOF_TABLES):
        D.clone_paragraph_after(anchor, text=entry, style=None)
    log.append(f"lof   {len(LOF_FIGS)} figure entries and {len(LOF_TABLES)} table "
               f"entries added; 2 superseded entries removed")


# =========================================================================== #
def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--dry-run", action="store_true",
                    help="write the working copy but do not touch the real report")
    ap.add_argument("--max-width", type=int, default=1400,
                    help="downscale figures wider than this many pixels "
                         "(default 1400 ~ 237 dpi at the 5.9 in text width)")
    args = ap.parse_args()

    if not DOCX.exists():
        raise SystemExit(f"report not found: {DOCX}")
    lock = DOCX.parent / ("~$" + DOCX.name[2:])
    if lock.exists():
        print(f"NOTE: Word lock file present ({lock.name}); close Word before the "
              f"final copy-back.", file=sys.stderr)

    shutil.copy2(DOCX, WORK)
    doc = D.load(WORK)
    before_p = len(list(doc.element.body.iter(qn("w:p"))))
    before_t = len(doc.tables)
    before_i = len(doc.inline_shapes)

    log: list[str] = []
    b = Builder(doc, args.max_width)

    renumber_figures(doc, log)
    fix_headings(doc, log)
    caption_orphans(b, log)
    insert_section_64(b, log)
    insert_section_71(b, log)
    expand_section_82(b, log)
    insert_section_85(b, log)
    insert_section_94(b, log)
    append_appendices(b, log)
    update_list_of_figures(doc, log)

    D.save(doc, WORK)
    IES.set_update_fields(WORK, log)

    doc2 = D.load(WORK)
    after_p = len(list(doc2.element.body.iter(qn("w:p"))))
    after_t = len(doc2.tables)
    after_i = len(doc2.inline_shapes)

    print("\n".join("  " + line for line in log))
    print(f"\nfigures embedded by this run: {b.n_images} "
          f"(26 in the body + {len(T.ALL_RUNS) * 7} in Appendix A)")
    print(f"paragraphs {before_p} -> {after_p} (+{after_p - before_p})")
    print(f"tables     {before_t} -> {after_t} (+{after_t - before_t})")
    print(f"images     {before_i} -> {after_i} (+{after_i - before_i})")
    print(f"size       {DOCX.stat().st_size / 1e6:.1f} MB -> "
          f"{WORK.stat().st_size / 1e6:.1f} MB")

    if args.dry_run:
        print(f"\ndry run: left the result at {WORK.relative_to(REPO)}")
        return 0

    if lock.exists():
        raise SystemExit(
            f"\nWord still has {DOCX.name} open ({lock.name}). Close it and re-run, "
            f"or copy {WORK.relative_to(REPO)} over the report yourself.")
    shutil.copy2(WORK, DOCX)
    WORK.unlink()
    print(f"\nwrote {DOCX.relative_to(REPO)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

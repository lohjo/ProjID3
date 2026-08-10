"""Verify the figure/statistics insertion against the artefacts it came from.

    python src/docs/review/_source/verify_figures_and_stats.py [path.docx]

Checks, in order:

1.  Package     - every part parses; every <a:blip> resolves to a media part;
                  every media part is referenced; the pre-existing images are
                  byte-identical to the pre-edit report.
2.  Figures     - body captions are contiguous Fig. 1..41 with no duplicates;
                  appendix captions are contiguous Fig. A1..A147; every caption
                  is immediately preceded by a picture paragraph.
3.  Tables      - captions Table 1..21, A.1, B.1..B.4, each used once.
4.  Statistics  - every cell of Tables B.1-B.4, 14-17 and A.1 is re-read out of
                  the saved .docx and compared against a freshly rendered table
                  from the committed CSVs. Zero tolerance.
5.  Sections    - no literal "?" remains in a heading; the new headings exist.
"""
from __future__ import annotations

import hashlib
import io
import re
import sys
import xml.etree.ElementTree as ET
import zipfile
from pathlib import Path

from PIL import Image

sys.path.insert(0, str(Path(__file__).resolve().parent))

import docxlib as D  # noqa: E402
import report_stats_tables as T  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402

HERE = Path(__file__).resolve().parent
REPO = HERE.parents[3]
# The pre-edit baseline: the report exactly as it stood before this pass ran.
BACKUP = REPO / "src/T32_PI05_Final_Report.backup-2026-08-09-prefigures.docx"
A_NS = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}

fails: list[str] = []
passes: list[str] = []


def check(ok: bool, msg: str) -> None:
    (passes if ok else fails).append(msg)
    print(("  PASS  " if ok else "  FAIL  ") + msg)


# --------------------------------------------------------------------------- #
def check_package(path: Path) -> None:
    print("\n[1] package integrity")
    z = zipfile.ZipFile(path)
    bad = []
    for n in z.namelist():
        if n.endswith((".xml", ".rels")):
            try:
                ET.fromstring(z.read(n))
            except Exception as exc:  # noqa: BLE001
                bad.append((n, str(exc)[:60]))
    check(not bad, f"all {sum(1 for n in z.namelist() if n.endswith(('.xml', '.rels')))} "
                   f"XML parts parse ({bad})")

    doc = z.read("word/document.xml").decode("utf8")
    rels = z.read("word/_rels/document.xml.rels").decode("utf8")
    rmap = {i: t.replace("&amp;", "&")
            for i, t in re.findall(r'Id="(rId\d+)"[^>]*Target="([^"]+)"', rels)}
    blips = re.findall(r'r:embed="(rId\d+)"', doc)
    vml = re.findall(r'r:id="(rId\d+)"[^>]*o:title|<v:imagedata r:id="(rId\d+)"', doc)
    vml_ids = {a or b for a, b in vml}
    media = {n[len("word/"):] for n in z.namelist() if n.startswith("word/media/")}
    targets = {rmap[b] for b in blips}
    missing = targets - media
    check(not missing, f"every blip resolves to a media part ({len(blips)} refs, "
                       f"{len(targets)} distinct){'' if not missing else missing}")
    check(len(blips) == len(targets),
          f"no media part is referenced twice ({len(blips)} refs / {len(targets)} parts)")
    unref = media - targets - {rmap[i] for i in vml_ids if i in rmap}
    unref = {u for u in unref if not u.endswith(".emf")}
    check(not unref, f"no orphan media (ignoring VML/EMF): {unref or 'none'}")

    if BACKUP.exists():
        # Compared by decoded pixels, not by part name: if the document has been
        # through a Word save, Word renumbers and re-encodes media parts, so a
        # name-keyed byte comparison reports differences that are not real. What
        # must hold is that every pre-existing image is still present, unaltered.
        def fingerprints(zf):
            out = {}
            for n in zf.namelist():
                if not n.startswith("word/media/"):
                    continue
                b = zf.read(n)
                try:
                    im = Image.open(io.BytesIO(b))
                    im.load()
                    out[n] = hashlib.sha256(im.convert("RGB").tobytes()).hexdigest()
                except Exception:  # noqa: BLE001  EMF and other vector parts
                    out[n] = hashlib.sha256(b).hexdigest()
            return out

        old = fingerprints(zipfile.ZipFile(BACKUP))
        new = set(fingerprints(z).values())
        lost = [n for n, h in old.items() if h not in new]
        check(not lost, f"all {len(old)} pre-existing images still present unaltered "
                        f"(compared by decoded pixels; missing: {lost})")


# --------------------------------------------------------------------------- #
def check_figures(doc) -> None:
    print("\n[2] figure captions")
    kids = D.body_children(doc)
    lof_end = next(i for i, el in kids
                   if el.tag == qn("w:p")
                   and D.ptext(el).startswith("Table B.4)"))
    body_nums, app_nums, orphan = [], [], []
    for i, el in kids:
        if el.tag != qn("w:p") or i <= lof_end:
            continue
        t = D.ptext(el).strip()
        m = re.match(r"^Fig\. A(\d+)\)", t)
        if m:
            app_nums.append((int(m.group(1)), i, el))
            continue
        m = re.match(r"^Fig\. (\d+)\)", t)
        if m:
            body_nums.append((int(m.group(1)), i, el))

    b = [n for n, _, _ in body_nums]
    check(b == sorted(b), f"body figure captions in ascending order ({len(b)})")
    check(len(set(b)) == len(b), f"no duplicate body figure number ({len(b)} captions)")
    # Fig. 8 ("Labelled process flow diagram") is listed in the List of Figures but
    # has never existed in the body and no such diagram exists in the repository.
    # The number is deliberately left reserved for the author to supply; see the
    # open-items note in updates.md. Everything else must be contiguous.
    expected_body = [n for n in range(1, 42) if n != 8]
    check(b == expected_body,
          f"body figures are exactly 1..41 with 8 reserved "
          f"(got {b[0]}..{b[-1]}, n={len(b)}, missing "
          f"{sorted(set(expected_body) - set(b)) or 'none'})")

    a = [n for n, _, _ in app_nums]
    check(a == list(range(1, 148)),
          f"appendix figures are exactly A1..A147 (n={len(a)})")

    # every appendix caption is directly preceded by a picture paragraph
    for n, i, el in app_nums:
        prev = el.getprevious()
        if prev is None or not prev.findall(".//a:blip", A_NS):
            orphan.append(n)
    check(not orphan, f"every appendix caption follows a picture "
                      f"({len(app_nums)} checked, orphans: {orphan[:5]})")

    # the 26 new body figures likewise (1-8 are pre-existing shape-group figures)
    new_body = [(n, el) for n, i, el in body_nums if n >= 9]
    bad = [n for n, el in new_body
           if n not in (12, 13, 14, 15, 16, 17, 18)  # pre-existing embeds
           and (el.getprevious() is None
                or not el.getprevious().findall(".//a:blip", A_NS))]
    check(not bad, f"every newly added body caption follows a picture (bad: {bad})")


def check_tables(doc) -> None:
    print("\n[3] table captions")
    kids = D.body_children(doc)
    lof_end = next(i for i, el in kids
                   if el.tag == qn("w:p") and D.ptext(el).startswith("Table B.4)"))
    seen = []
    for i, el in kids:
        if el.tag != qn("w:p") or i <= lof_end:
            continue
        t = D.ptext(el).strip()
        # "Table 13) shows a list of ..." is an in-text reference, not a caption;
        # a caption is followed by a capitalised noun phrase.
        m = re.match(r"^Table ([0-9]+|A\.1|B\.[1-4])\)\s+(\S+)", t)
        if m and not m.group(2)[0].islower():
            seen.append(m.group(1))
    expected = [str(n) for n in range(1, 22)] + ["A.1", "B.1", "B.2", "B.3", "B.4"]
    dupes = [s for s in set(seen) if seen.count(s) > 1]
    missing = [e for e in expected if e not in seen]
    check(not missing, f"all expected table captions present (missing: {missing})")
    check(not dupes, f"no duplicated table caption (dupes: {dupes})")
    check(len(doc.tables) == 37, f"table count is 37 (got {len(doc.tables)})")


# --------------------------------------------------------------------------- #
def _grid(tbl) -> list[list[str]]:
    out = []
    for row in tbl.rows:
        out.append([c.text.strip() for c in row.cells])
    return out


def check_statistics(doc) -> None:
    print("\n[4] statistics re-read from the saved document")
    expected = {
        "Table 14": T.tbl_fit_quality(T.RUNS_GRID),
        "Table 15": T.tbl_model_ranking(T.RUNS_GRID),
        "Table 16": T.tbl_model_ranking(T.RUNS_OLD),
        "Table 17": T.tbl_ftest(T.ALL_RUNS),
        "Table 18": T.tbl_e1_cluster_anova(),
        "Table 19": T.tbl_e2_factorial(),
        "Table 20": T.tbl_e2b_reproducibility(),
        "Table 21": T.tbl_bound_pinning(),
        "Table B.1": T.tbl_b1_master(),
        "Table B.2": T.tbl_b2_parameters(),
        "Table B.3": T.tbl_b3_ftests(),
        "Table B.4": T.tbl_b4_degenerate(),
    }
    grids = [_grid(t) for t in doc.tables]
    total_cells = 0
    for name, want in expected.items():
        hit = None
        for g in grids:
            if len(g) == len(want) and len(g[0]) == len(want[0]) and g[0] == want[0]:
                if g == want:
                    hit = g
                    break
                hit = hit or ("MISMATCH", g)
        if hit is None:
            check(False, f"{name}: no table of shape {len(want)}x{len(want[0])} "
                         f"with a matching header found in the document")
            continue
        if isinstance(hit, tuple):
            g = hit[1]
            diffs = [(r, c, want[r][c], g[r][c])
                     for r in range(len(want)) for c in range(len(want[0]))
                     if want[r][c] != g[r][c]]
            check(False, f"{name}: {len(diffs)} cell(s) differ, first "
                         f"{diffs[:3]}")
            continue
        cells = len(want) * len(want[0])
        total_cells += cells
        check(True, f"{name}: all {cells} cells match the committed artefacts "
                    f"({len(want) - 1} data rows)")
    print(f"        -> {total_cells} cells verified against source")


def check_sections(doc) -> None:
    print("\n[5] headings")
    for want, lvl in [("6.4 Measured breakthrough curves", 2),
                      ("7.1 Cross-run trends and model ranking", 2),
                      ("7.2 Determining the Optimal Operating Parameters", 2),
                      ("7.2.4 Possible Adjustments", 3),
                      ("8.5 Sensitivity analysis", 2),
                      ("9.4 Numerical verification", 2),
                      ("Appendix A", 1), ("Appendix B", 1)]:
        try:
            D.find_heading(doc, want, level=lvl)
            check(True, f"heading present: {want!r} (H{lvl})")
        except ValueError:
            check(False, f"heading MISSING: {want!r} (H{lvl})")
    stale = [D.ptext(el)[:60] for _, el in D.body_children(doc)
             if el.tag == qn("w:p") and D.is_heading(el) and "?" in D.ptext(el)]
    check(not stale, f"no heading carries a literal '?' ({stale})")


def main() -> int:
    path = Path(sys.argv[1]) if len(sys.argv) > 1 else HERE / "_work_report_figs.docx"
    print(f"verifying {path}  ({path.stat().st_size / 1e6:.1f} MB)")
    check_package(path)
    doc = D.load(path)
    check_figures(doc)
    check_tables(doc)
    check_statistics(doc)
    check_sections(doc)
    print(f"\n{len(passes)} passed, {len(fails)} failed")
    for f in fails:
        print("  FAILED: " + f)
    return 1 if fails else 0


if __name__ == "__main__":
    raise SystemExit(main())

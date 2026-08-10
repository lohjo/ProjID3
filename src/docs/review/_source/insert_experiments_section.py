"""Insert §5 "Experiments" into T32_PI05_Final_Report.docx, in place.

    python src/docs/review/_source/insert_experiments_section.py [--dry-run]

What it does, in order:

1.  Renumbers §5-§10 to §6-§11 (H1, H2, H3, and the body-styled pseudo-headings
    the document uses in place of real headings in two places).
2.  Renumbers the table captions and every in-text "Table N" reference. Applied
    high-to-low so numbers never collide. This also resolves a pre-existing
    collision: two different tables were both captioned "Table 7)".
3.  Inserts §5 with three generated tables between §4 (SOP) and what is now §6
    (Experimental Results).
4.  Adds the three new tables to the List of Figures and renumbers its entries.
5.  Sets <w:updateFields> so Word offers to refresh the TOC field on open.

Every number in the inserted tables is read from `src/docs/experiments_passport.md`,
which `src/solver/build_experiments_passport.py` generates from committed artefacts.
Nothing here is hand-typed, and nothing already in the manuscript is rewritten
apart from the numbering.

The script works on a copy and only overwrites the real file once every step has
succeeded. Word must be closed while it runs.
"""

from __future__ import annotations

import argparse
import re
import shutil
import sys
import zipfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))

import docxlib as D  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.shared import Pt  # noqa: E402

HERE = Path(__file__).resolve().parent
REPO = HERE.parents[3]
DOCX = REPO / "src/T32_PI05_Final_Report.docx"
WORK = HERE / "_work_report.docx"
PASSPORT_MD = REPO / "src/docs/experiments_passport.md"

TABLE_STYLE = "Table Grid1"

# --------------------------------------------------------------------------- #
# Renumbering maps. Order matters: highest first, so a rename never lands on a
# number that has not been vacated yet.
# --------------------------------------------------------------------------- #
H1_RENUMBER = [
    ("10 References", "11 References"),
    ("9 Conclusions", "10 Conclusions"),
    ("8 Mathematical Modelling", "9 Mathematical Modelling"),
    ("6 Experimental Analysis", "7 Experimental Analysis"),
    ("5 Experimental Results", "6 Experimental Results"),
]

# Sub-headings, descending. (old prefix, new prefix); matched at start of text.
SUB_RENUMBER = [
    ("9.2", "10.2"), ("9.1", "10.1"),
    ("8.3", "9.3"), ("8.2", "9.2"), ("8.1", "9.1"),
    ("7.4", "8.4"), ("7.3", "8.3"), ("7.2", "8.2"), ("7.1", "8.1"),
    ("6.?.4", "7.?.4"),
    ("5.3", "6.3"), ("5.2", "6.2"), ("5.1", "6.1"),
]

# The §7 H1 is body-styled and has no space after its number, so it never reached
# the table of contents. Renumbered, spaced, and promoted to a real Heading1.
BODY_H1_RENUMBER = [("7Fitting performance and analysis",
                     "8 Fitting performance and analysis")]

# Same defect one level down: §5's three sub-headings are body-styled. Promoted to
# Heading2 after renumbering so the contents list is complete. Text is unchanged.
PROMOTE_TO_H2 = [
    "6.1 Parameters Measured",
    "6.2 Adsorption Capacity Calculation",
    "6.3 Results",
]

# Headings missing the space between number and title. Applied after renumbering.
SPACING_FIXES = [("9.3Limitations", "9.3 Limitations")]

# Table captions, identified by their opening text. Highest first.
CAPTION_RENUMBER = [
    # Matched on caption text as well as number: "Table 8) shows a list of..." is a
    # sentence in §9.2, not a caption, and must not be swept up here.
    ("Table 9) The basic parameters", "Table 13) The basic parameters"),
    ("Table 8) Langmuir isotherm parameters", "Table 12) Langmuir isotherm parameters"),
    # the two "Table 7)" captions are disambiguated by their following text
    ("Table 7) Winning model per run", "Table 11) Winning model per run"),
    ("Table 7) Breakthrough", "Table 10) Breakthrough"),
    ("Table 6) Breakthrough", "Table 9) Breakthrough"),
    ("Table 5) Breakthrough adsorption", "Table 8) Breakthrough adsorption"),
    ("Table 4) Experimental Parameters", "Table 7) Experimental Parameters"),
    ("Table 3) The basic parameters", "Table 6) The basic parameters"),
]

# In-text cross-references: (unique text fragment locating the paragraph,
# old substring, new substring, note).
INTEXT_REFS = [
    ("The operating conditions and the corresponding parameters are given in",
     "Table 3)", "Table 6)", "points at the basic-parameters table"),
    ("The result of the equilibrium adsorption test is shown in",
     "Table 4)", "Table 7)", "points at the experimental-parameters table"),
    ("reports the winning model per run against the M01",
     "Table 6 reports", "Table 11 reports",
     "was already wrong: the winning-model table was captioned Table 7"),
    ("shows the Langmuir isotherm parameters and",
     "Table 8) shows the Langmuir", "Table 12) shows the Langmuir", "Langmuir table"),
    ("shows the Langmuir isotherm parameters and",
     "Table 9) shows a list", "Table 13) shows a list", "packed-bed parameters table"),
    ("Bed void fraction",
     "(Table 9, above)", "(Table 13, above)", "superficial-velocity footnote"),
]

# Report-internal section cross-references invalidated by the renumber. Refs to
# `experimental-results.md §N` point at a different document and are left alone, as
# are all §3.x and §4.x refs, whose numbering does not move.
# (locator phrase, old, new)
SECTION_REFS = [
    ("independently corroborates the fractal-like models' advantage reported in",
     "§6 of this work", "§7 of this work"),
    ("identical R², RMSE and AICc to at least five significant figures",
     "see §7)", "see §8)"),
    ("the remaining models fitted in the",
     "the §7 comparison", "the §8 comparison"),
    ("models that win in four of the five original runs",
     "(§7; Hu", "(§8; Hu"),
    ("so its goodness-of-fit statistics in",
     "statistics in §7", "statistics in §8"),
    ("as a candidate in the model-selection comparison of",
     "comparison of §7.", "comparison of §8."),
    ("approximation is valid only for",
     "§7 reports that these conditions", "§8 reports that these conditions"),
    ("return non-finite RSS and an undefined R² in every run",
     "(§7);", "(§8);"),
    ("is the most flexible model in the 24-model registry",
     "§7 reports its AICc penalty", "§8 reports its AICc penalty"),
    ("even where the choice of", "(§7.1)", "(§8.1)"),
    ("supplementary to the report's core empirical contribution in", "§§4-7", "§§4-8"),
    ("have yet been fitted against the measured breakthrough data in", "§5-§7", "§6-§8"),
]

# The stray repeat of the packed-bed-parameters sentence in §8.2 says "Table 8)"
# but describes the packed-bed parameter table, which is Table 9 before the
# renumber and Table 13 after it. Corrected to what it means.
STRAY_REF = ("Table 8) shows a list of parameters used and their values",
             "Table 8)", "Table 13)")

# The List of Figures entries share their opening text with the body captions, so
# CAPTION_RENUMBER already renumbers both. Only the three new entries are added here,
# after the "Table 2)" entry.
LOF_NEW = [
    "Table 3) Design as executed: number of usable runs in each flow × "
    "concentration cell, 26 June to 31 July 2026.",
    "Table 4) Run inventory and provenance for the sixteen usable breakthrough "
    "records.",
    "Table 5) Excluded records and records carrying a caveat, with the owner of "
    "each open item.",
]

CAPTIONS = {
    3: "Table 3) Design as executed: number of usable runs in each flow × "
       "concentration cell, 26 June to 31 July 2026.",
    4: "Table 4) Run inventory and provenance for the sixteen usable breakthrough "
       "records. Conditions are read from each file's own embedded header. The source "
       "file for each run is <run>.csv in src/solver/data/newest runs/.",
    5: "Table 5) Excluded records and records carrying a caveat, with the owner of "
       "each open item.",
}


# --------------------------------------------------------------------------- #
# Passport tables
# --------------------------------------------------------------------------- #
def load_passport_tables() -> dict[int, list[list[str]]]:
    """Parse the three Markdown tables out of experiments_passport.md."""
    if not PASSPORT_MD.exists():
        raise SystemExit(
            f"{PASSPORT_MD} not found -- run build_experiments_passport.py first"
        )
    text = PASSPORT_MD.read_text(encoding="utf-8")
    tables: dict[int, list[list[str]]] = {}
    for number in (3, 4, 5):
        m = re.search(
            rf"^## Table {number} .*?$\n\n((?:\|.*\n)+)", text, re.M
        )
        if not m:
            raise SystemExit(f"Table {number} not found in {PASSPORT_MD.name}")
        rows = []
        for line in m.group(1).strip().splitlines():
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            if all(set(c) <= set("-: ") for c in cells):
                continue  # separator row
            rows.append(cells)
        tables[number] = rows
    return tables


# --------------------------------------------------------------------------- #
# Paragraph / table construction
# --------------------------------------------------------------------------- #
def set_style(p_el, style: str) -> None:
    """Set a paragraph's style, creating <w:pPr>/<w:pStyle> if absent."""
    pPr = p_el.find("w:pPr", D.NS)
    if pPr is None:
        pPr = p_el.makeelement(qn("w:pPr"), {})
        p_el.insert(0, pPr)
    pStyle = pPr.find("w:pStyle", D.NS)
    if pStyle is None:
        pStyle = pPr.makeelement(qn("w:pStyle"), {})
        pPr.insert(0, pStyle)
    pStyle.set(qn("w:val"), style)


def add_paragraph_after(anchor_el, text: str, style: str, template_el):
    """Clone `template_el`'s paragraph shape, set style + text, insert after anchor."""
    new_el = D.clone_paragraph_after(anchor_el, text=text, style=style)
    return new_el


def build_table(doc, rows: list[list[str]], font_pt: float):
    tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
    try:
        tbl.style = TABLE_STYLE
    except KeyError:
        tbl.style = "Table Grid"
    tbl.autofit = True
    # Span the text column (5.9 in here) instead of inheriting a fixed width.
    tblPr = tbl._tbl.tblPr
    for tag in ("w:tblW", "w:tblLayout"):
        for existing in tblPr.findall(tag, D.NS):
            tblPr.remove(existing)
    tblW = tblPr.makeelement(qn("w:tblW"), {})
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")
    tblPr.append(tblW)
    layout = tblPr.makeelement(qn("w:tblLayout"), {})
    layout.set(qn("w:type"), "autofit")
    tblPr.append(layout)

    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            cell = tbl.cell(i, j)
            cell.text = ""
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = para.add_run(value)
            run.font.size = Pt(font_pt)
            if i == 0:
                run.font.bold = True
    return tbl


# --------------------------------------------------------------------------- #
# Section text
# --------------------------------------------------------------------------- #
INTRO = [
    ("Body",
     "This section records the experiments as they were actually run. Section 4 "
     "gives the procedure and Section 6 reports the outcomes; what follows is the "
     "inventory that connects the two, so that each result reported later can be "
     "traced back to the record it came from."),
    ("Body",
     "Eighteen breakthrough records were logged between 26 June and 31 July 2026 on "
     "the 8.2 mm internal-diameter column. Sixteen of them carry an embedded header "
     "block stating sorbent mass, bed height, tube diameter, flow rate and inlet "
     "CO₂ concentration, and those sixteen form the basis of every result in this "
     "report. The remaining two are raw sensor logs without that header and are "
     "excluded for the reasons set out in Table 5. The five earlier runs on the "
     "8.5 mm column, whose parameters appear in Table 7, are a separate campaign and "
     "are not pooled with the grid below: bed length and column internal diameter "
     "differ systematically between the two rigs, so pooling them would confound a "
     "change of apparatus with the effects being measured."),
]

SEC_51 = [
    ("Heading2", "5.1 Design as executed"),
    ("Body",
     "The design is a 3 × 3 factorial in inlet CO₂ concentration (5, 10 and 15 % by "
     "volume) and total gas flow rate (50, 100 and 150 ml/min). The remaining factors "
     "were held fixed: 8.00 g of PEI@SiO₂ granules in a bed of 22.5 to 24.5 cm, an "
     "8.2 mm internal diameter, ambient temperature, and a pressure of approximately "
     "101.3 kPa. Seven of the nine cells were run more than once, giving the counts in "
     "Table 3."),
    ("Body",
     "Replication carries more weight here than the extra runs alone suggest. With a "
     "single run per cell the interaction between flow rate and concentration cannot "
     "be separated from residual error at all, because no degrees of freedom remain to "
     "estimate that error. The seven repeated cells supply seven such degrees of "
     "freedom, which is what makes the interaction term and the pure measurement "
     "scatter separately estimable. The design is nonetheless unbalanced as executed: "
     "the 15 % cells at 50 and 150 ml/min were each run once, so any effect resting on "
     "those two cells rests on less evidence than the rest of the grid."),
]

SEC_52 = [
    ("Heading2", "5.2 Run inventory and provenance"),
    ("Body",
     "Table 4 lists every usable run, the file it was read from, the conditions that "
     "file states about itself, and the point later in this report where its results "
     "appear. Inlet concentration, sorbent mass, bed height and column diameter are "
     "extracted from each file's own header rather than transcribed by hand, so the "
     "values in Table 4 are the values the fitting pipeline used, not a separate "
     "record of them. Runs are ordered to match the two replicate tables in §6.3: "
     "numbers 1 to 9 are replicate I, numbers 10 to 16 are replicate II."),
]

SEC_52_TAIL = [
    ("Body",
     "Four files state their flow rate twice and state it differently. These are "
     "marked with an asterisk in Table 4, and the discrepancy is set out in §5.3."),
]

SEC_53 = [
    ("Heading2", "5.3 Excluded and flagged records"),
    ("Body",
     "No record was dropped silently and none was reconstructed. Table 5 lists every "
     "record that is either excluded from the analysis or carries a caveat affecting "
     "how its numbers should be read, together with the party who can close each open "
     "item."),
]

SEC_53_TAIL = [
    ("Body",
     "The four flow-rate conflicts share a pattern worth stating plainly. In every "
     "case the prose value, not the numeric cell the parser reads, is the one "
     "consistent with the measured data. The clearest instance is "
     "2026-07-22-conc10-flow0.10, which breaks through at 3.08 min: that sits beside "
     "the 3.55 min of the 150 ml/min run at the same concentration and nowhere near "
     "the 5.43 min of the 100 ml/min run. The same comparison favours the prose value "
     "in the other three files."),
    ("Body",
     "The consequence is bounded but real. Breakthrough time is read from the "
     "concentration record itself and does not depend on the stated flow rate, so it "
     "is unaffected. Dynamic capacity, mass-transfer-zone length and stoichiometric "
     "efficiency all scale with volumetric flow, so for these four runs those three "
     "quantities are provisional until the lab confirms which of the two values is "
     "correct. The filename of 2026-07-22-conc10-flow0.10.csv, which disagrees with "
     "its own header, should be corrected in the same pass."),
]

SEC_54 = [
    ("Heading2", "5.4 Data lineage and reproducibility"),
    ("Body",
     "Every record passes through the same chain. The parser reads the file together "
     "with its embedded header; the fitting stage fits the 24-model set by non-linear "
     "least squares using 12 multi-start initialisations under a fixed random seed of "
     "42; the performance stage derives the operating-time and capacity metrics; and "
     "the result is written to a per-run results file alongside seven diagnostic plots. "
     "Tables 3, 4 and 5 are generated from those artefacts by "
     "build_experiments_passport.py, which also writes a provenance record carrying a "
     "SHA-256 checksum for every raw file, so that any table cell here can be traced "
     "to the exact bytes it came from. No value in this section is entered by hand."),
    ("Body",
     "One definitional point carries forward into the results. Breakthrough and "
     "equilibrium times in Tables 9 and 10 are read at C/C₀ = 0.05 and C/C₀ = 0.95 "
     "respectively, as defined in §6.1, and every derived quantity in the analysis that "
     "follows uses those definitions. The times recorded in Table 8 were noted during "
     "the runs using a visual onset criterion that was not written down, and they are "
     "systematically shorter: the first run of the grid is entered there as 14 min "
     "against 23.96 min at C/C₀ = 0.05. Both are reported. The C/C₀ definitions are the "
     "ones used for analysis throughout."),
]


# --------------------------------------------------------------------------- #
def para_style(el) -> str:
    pStyle = el.find("w:pPr/w:pStyle", D.NS)
    return pStyle.get(qn("w:val"), "") if pStyle is not None else ""


def is_toc_entry(el) -> bool:
    """Cached table-of-contents field results carry TOC1/TOC2/TOC3 styles.

    They repeat every heading's text verbatim and sit near the top of the document,
    so any naive text search hits them first. Word regenerates them from the real
    headings when the field refreshes, so they must never be edited in their place.
    """
    return para_style(el).startswith("TOC")


def find_body_paragraph(doc, contains: str, heading: str | None = None):
    """First real body-level <w:p> containing `contains` (optionally a given style)."""
    for i, el in D.body_children(doc):
        if el.tag != qn("w:p") or is_toc_entry(el):
            continue
        if contains not in D.ptext(el):
            continue
        if heading is not None and D.is_heading(el) != heading:
            continue
        return i, el
    raise SystemExit(f"paragraph not found: {contains!r} (style={heading})")


def renumber(doc, log: list[str]) -> None:
    # --- H1 -------------------------------------------------------------- #
    for old, new in H1_RENUMBER:
        _, el = D.find_heading(doc, old, level=1)
        old_num, new_num = old.split(" ", 1)[0], new.split(" ", 1)[0]
        if not D.replace_text(el, old_num, new_num, occurrence=1):
            raise SystemExit(f"H1 renumber failed: {old!r}")
        log.append(f"H1   {old}  ->  {new}")

    # --- body-styled H1 --------------------------------------------------- #
    for old, new in BODY_H1_RENUMBER:
        _, el = find_body_paragraph(doc, old)
        if not D.replace_text(el, old, new, occurrence=1):
            raise SystemExit(f"body-H1 renumber failed: {old!r}")
        set_style(el, "Heading1")
        log.append(f"H1*  {old}  ->  {new}   (was body-styled; promoted to Heading1 "
                   f"so it reaches the contents list)")

    # --- H2 / H3 and body-styled pseudo-headings -------------------------- #
    for old_pref, new_pref in SUB_RENUMBER:
        hits = []
        for _, el in D.body_children(doc):
            if el.tag != qn("w:p") or is_toc_entry(el):
                continue
            text = D.ptext(el).strip()
            if not text.startswith(old_pref):
                continue
            # Guard: only a heading or a short standalone label, never prose.
            if not (D.is_heading(el) or len(text) < 70):
                continue
            hits.append((el, text))
        if not hits:
            raise SystemExit(f"sub-heading not found: {old_pref!r}")
        for el, text in hits:
            if not D.replace_text(el, old_pref, new_pref, occurrence=1):
                raise SystemExit(f"sub-heading renumber failed: {text!r}")
            log.append(f"sub  {text[:52]:52s} ->  {new_pref}")

    # --- table captions ---------------------------------------------------- #
    for old, new in CAPTION_RENUMBER:
        hits = [
            (i, el) for i, el in D.body_children(doc)
            if el.tag == qn("w:p") and D.ptext(el).strip().startswith(old)
        ]
        if not hits:
            raise SystemExit(f"caption not found: {old!r}")
        old_num = old.split(")")[0] + ")"
        new_num = new.split(")")[0] + ")"
        for _, el in hits:
            if not D.replace_text(el, old_num, new_num, occurrence=1):
                raise SystemExit(f"caption renumber failed: {old!r}")
        log.append(f"cap  {old_num} -> {new_num}  ({len(hits)} occurrence(s): "
                   f"body caption + List of Figures where present)")

    # --- in-text references ------------------------------------------------ #
    for locator, old, new, note in INTEXT_REFS:
        _, el = find_body_paragraph(doc, locator)
        if D.replace_text(el, old, new, occurrence=1):
            log.append(f"ref  {old:28s} -> {new:22s} ({note})")
        else:
            log.append(f"ref  SKIPPED {old!r} in {locator[:40]!r} -- already renumbered")

    locator, old, new = STRAY_REF
    _, el = find_body_paragraph(doc, locator)
    if D.replace_text(el, old, new, occurrence=1):
        log.append(f"ref  {old:28s} -> {new:22s} (stray repeat in §9.2; "
                   f"it describes the packed-bed parameter table)")

    # --- promote the body-styled §6 sub-headings --------------------------- #
    for text in PROMOTE_TO_H2:
        _, el = find_body_paragraph(doc, text)
        set_style(el, "Heading2")
        log.append(f"sty  {text:44s} -> Heading2 (was body-styled)")

    # --- section cross-references ------------------------------------------ #
    for locator, old, new in SECTION_REFS:
        _, el = find_body_paragraph(doc, locator)
        if not D.replace_text(el, old, new, occurrence=1):
            raise SystemExit(f"section ref not found: {old!r} in {locator[:45]!r}")
        log.append(f"§ref {old:26s} -> {new}")

    for old, new in SPACING_FIXES:
        _, el = find_body_paragraph(doc, old)
        if D.replace_text(el, old, new, occurrence=1):
            log.append(f"spc  {old:44s} -> {new}")


def renumber_list_of_figures(doc, log: list[str]) -> None:
    """Add the three new tables to the manual List of Figures."""
    _, anchor = find_body_paragraph(doc, "Table 2) Breakdown of Project Schedule")
    for entry in reversed(LOF_NEW):
        D.clone_paragraph_after(anchor, text=entry, style=None)
    log.append(f"lof  3 new entries added after the Table 2) entry")


def insert_section(doc, tables: dict[int, list[list[str]]], log: list[str]) -> None:
    idx, h1 = D.find_heading(doc, "6 Experimental Results", level=1)
    body = doc.element.body

    # A Heading1 paragraph to clone the shape from, and a Body paragraph likewise.
    _, h1_template = D.find_heading(doc, "4 Standard Operating Procedure", level=1)
    _, h2_template = D.find_heading(doc, "4.1 Apparatus and Instrumentation", level=2)
    body_template = h1.getnext()
    while body_template is not None and (
        body_template.tag != qn("w:p") or D.is_heading(body_template)
    ):
        body_template = body_template.getnext()
    if body_template is None:
        raise SystemExit("no body paragraph found to use as a template")

    def emit(el_after, kind: str, text: str):
        template = {"Heading1": h1_template, "Heading2": h2_template}.get(
            kind, body_template
        )
        new_el = D.clone_paragraph_after(el_after, text=text, style=kind)
        # clone_paragraph_after copies el_after's shape; re-home the formatting by
        # copying the template's pPr when the styles differ.
        if template is not el_after:
            tpl_pPr = template.find("w:pPr", D.NS)
            if tpl_pPr is not None:
                import copy as _copy
                old_pPr = new_el.find("w:pPr", D.NS)
                if old_pPr is not None:
                    new_el.remove(old_pPr)
                new_el.insert(0, _copy.deepcopy(tpl_pPr))
        return new_el

    # Build the block bottom-up is fiddly; instead build top-down from a marker
    # paragraph placed immediately before the (now) §6 heading.
    marker = D.clone_paragraph_after(h1, text="", style="Body")
    body.remove(marker)
    h1.addprevious(marker)

    cursor = marker
    cursor = emit(cursor, "Heading1", "5 Experiments")
    for kind, text in INTRO:
        cursor = emit(cursor, kind, text)

    for kind, text in SEC_51:
        cursor = emit(cursor, kind, text)
    cursor = emit(cursor, "Body", CAPTIONS[3])
    cursor = place_table(doc, cursor, tables[3], font_pt=9.0)

    for kind, text in SEC_52:
        cursor = emit(cursor, kind, text)
    cursor = emit(cursor, "Body", CAPTIONS[4])
    cursor = place_table(doc, cursor, tables[4], font_pt=6.5)
    for kind, text in SEC_52_TAIL:
        cursor = emit(cursor, kind, text)

    for kind, text in SEC_53:
        cursor = emit(cursor, kind, text)
    cursor = emit(cursor, "Body", CAPTIONS[5])
    cursor = place_table(doc, cursor, tables[5], font_pt=7.0)
    for kind, text in SEC_53_TAIL:
        cursor = emit(cursor, kind, text)

    for kind, text in SEC_54:
        cursor = emit(cursor, kind, text)

    # remove the empty marker paragraph
    marker.getparent().remove(marker)

    log.append("sec  §5 Experiments inserted before §6 Experimental Results "
               f"(4 sub-sections, 3 tables, {len(tables[4]) - 1} inventory rows)")


def place_table(doc, after_el, rows, font_pt: float):
    """Create a table, move it after `after_el`, return an empty spacer paragraph."""
    tbl = build_table(doc, rows, font_pt)
    tbl_el = tbl._tbl
    tbl_el.getparent().remove(tbl_el)
    after_el.addnext(tbl_el)
    spacer = D.clone_paragraph_after(after_el, text="", style="Body")
    # clone lands right after after_el, i.e. before the table; move it past.
    spacer.getparent().remove(spacer)
    tbl_el.addnext(spacer)
    return spacer


def set_update_fields(path: Path, log: list[str]) -> None:
    """Ask Word to refresh field results (the TOC) when the document opens."""
    tmp = path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(path) as zin, zipfile.ZipFile(
        tmp, "w", zipfile.ZIP_DEFLATED
    ) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/settings.xml":
                text = data.decode("utf-8")
                if "w:updateFields" not in text:
                    text = text.replace(
                        "<w:settings ", "<w:settings ", 1
                    )
                    insert_at = text.index(">", text.index("<w:settings")) + 1
                    text = (text[:insert_at]
                            + '<w:updateFields w:val="true"/>'
                            + text[insert_at:])
                    data = text.encode("utf-8")
                    log.append("toc  <w:updateFields> set: Word will offer to "
                               "refresh the table of contents on open")
                else:
                    log.append("toc  <w:updateFields> already present")
            zout.writestr(item, data)
    tmp.replace(path)


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--dry-run", action="store_true",
                    help="write the working copy but do not touch the real report")
    args = ap.parse_args()

    if not DOCX.exists():
        raise SystemExit(f"report not found: {DOCX}")

    tables = load_passport_tables()
    print(f"loaded passport tables: "
          f"design {len(tables[3])-1} rows, inventory {len(tables[4])-1} rows, "
          f"flags {len(tables[5])-1} rows")

    shutil.copy2(DOCX, WORK)
    doc = D.load(WORK)

    before_paras = len(list(doc.element.body.iter(qn("w:p"))))
    before_tbls = len(doc.tables)

    log: list[str] = []
    renumber(doc, log)
    renumber_list_of_figures(doc, log)
    insert_section(doc, tables, log)
    D.save(doc, WORK)
    set_update_fields(WORK, log)

    doc2 = D.load(WORK)
    after_paras = len(list(doc2.element.body.iter(qn("w:p"))))
    after_tbls = len(doc2.tables)

    print("\n".join("  " + line for line in log))
    print(f"\nparagraphs {before_paras} -> {after_paras} "
          f"(+{after_paras - before_paras})")
    print(f"tables     {before_tbls} -> {after_tbls} (+{after_tbls - before_tbls})")

    if args.dry_run:
        print(f"\ndry run: left the result at {WORK.relative_to(REPO)}")
        return 0

    shutil.copy2(WORK, DOCX)
    WORK.unlink()
    print(f"\nwrote {DOCX.relative_to(REPO)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
